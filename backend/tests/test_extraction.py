"""Integration tests for text and table extraction (Task 12)."""

import os
from pathlib import Path

import pytest
import structlog

from app.modules.ingestion.extraction import (
    extract_docx_content,
    extract_pdf_content,
    strip_headers_footers,
)

logger = structlog.get_logger(__name__)

# Path to the test fixture
FIXTURE_PATH = Path(__file__).parent.parent.parent / "plan" / "fixtures" / "RFP" / "tele-manas-rfp-sample.pdf"


def _check_fixture() -> None:
    """Check if fixture exists, fail in CI if missing, skip otherwise."""
    if not FIXTURE_PATH.exists():
        if os.environ.get("CI"):
            pytest.fail(f"Required fixture missing in CI: {FIXTURE_PATH}")
        pytest.skip(f"Fixture not found at {FIXTURE_PATH} (skipping outside CI)")


@pytest.mark.integration
def test_pdf_extraction_tele_manas_fixture():
    """Test PDF extraction against the Tele-MANAS RFP fixture.

    This validates that:
    1. Extraction produces blocks across all ~68 pages
    2. Headers/footers are stripped (repeated "Request for Proposal" boilerplate)
    3. Tables are detected (Eligibility Criteria, Rate Card)
    4. Tables have correct row/column counts
    """
    _check_fixture()

    with open(FIXTURE_PATH, "rb") as f:
        file_data = f.read()

    # Extract content
    blocks = extract_pdf_content(file_data)

    # Basic validation: should have blocks from all pages
    pages_with_blocks = set(b["page"] for b in blocks)
    assert len(pages_with_blocks) >= 60, f"Expected blocks from ~68 pages, got {len(pages_with_blocks)}"

    # Check for table blocks
    table_blocks = [b for b in blocks if b["type"] == "table"]
    assert len(table_blocks) >= 2, f"Expected at least 2 tables (Eligibility Criteria, Rate Card), got {len(table_blocks)}"

    # Validate Eligibility Criteria table (Annexure 3) - should have rows with eligibility criteria
    # Note: pdfplumber may split merged cells, so we look for the table with eligibility headers
    eligibility_table = None
    for tb in table_blocks:
        content = tb["content"]
        if len(content) >= 7 and len(content[0]) >= 3:
            # Check if first row contains eligibility-like headers
            first_row = " ".join(str(c).lower() for c in content[0] if c)
            if "requirement" in first_row or "criteria" in first_row or "supporting" in first_row:
                eligibility_table = tb
                break

    assert eligibility_table is not None, "Eligibility Criteria table not found"
    assert len(eligibility_table["content"]) >= 7, f"Eligibility table should have at least 7 rows, got {len(eligibility_table['content'])}"
    assert len(eligibility_table["content"][0]) >= 3, "Eligibility table should have at least 3 columns"

    # Validate Rate Card table (Annexure 9) - should have role names
    rate_card_table = None
    for tb in table_blocks:
        content = tb["content"]
        if len(content) >= 9 and len(content[0]) >= 3:
            # Check for rate card roles
            all_text = " ".join(str(c).lower() for row in content for c in row if c)
            if "project manager" in all_text and "programmer" in all_text and "architect" in all_text:
                rate_card_table = tb
                break

    assert rate_card_table is not None, "Rate Card table (Annexure 9) not found"
    assert len(rate_card_table["content"]) >= 9, f"Rate card should have at least 9 roles, got {len(rate_card_table['content'])}"

    # Verify headers/footers are stripped
    # The repeated "Request for Proposal" / title / page-number boilerplate should not appear
    # as standalone repeated blocks in the output
    header_footer_content = []
    for block in blocks:
        if block["type"] == "text":
            content_lower = block["content"].lower()
            if ("request for proposal" in content_lower and
                "iiitb" in content_lower and
                len(block["content"]) < 200):
                header_footer_content.append(block["content"])

    # Should not have many repeated header/footer blocks
    # (The stripper should have removed most of them)
    assert len(header_footer_content) < 5, (
        f"Too many header/footer blocks remain ({len(header_footer_content)}). "
        "Header/footer stripping may not be working correctly."
    )

    # Verify real content is retained
    real_content_blocks = [b for b in blocks if b["type"] in ("text", "heading")]
    total_chars = sum(len(b["content"]) for b in real_content_blocks)
    assert total_chars > 50000, f"Expected substantial content extraction, got {total_chars} chars"

    logger.info(
        "extraction_test_passed",
        total_blocks=len(blocks),
        pages=len(pages_with_blocks),
        tables=len(table_blocks),
        text_blocks=len(real_content_blocks),
    )


@pytest.mark.integration
def test_header_footer_stripper_removes_repeated_blocks():
    """Test that header_footer_stripper correctly identifies and removes repeated blocks."""
    # Create mock blocks with repeated header/footer
    # Use realistic page dimensions: letter size 612x792
    # Header at top (y=0-30), footer at bottom (y=760-792), content in middle (y=150-600)
    blocks = [
        {"type": "text", "page": 1, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 1, "content": "Real content on page 1", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 1, "content": "Page 1", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 2, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 2, "content": "Real content on page 2", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 2, "content": "Page 2", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 3, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 3, "content": "Real content on page 3", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 3, "content": "Page 3", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 4, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 4, "content": "Real content on page 4", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 4, "content": "Page 4", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 5, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 5, "content": "Real content on page 5", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 5, "content": "Page 5", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 6, "content": "Request for Proposal", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 6, "content": "Real content on page 6", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 6, "content": "Page 6", "bbox": [0, 760, 612, 792]},
    ]

    stripped = strip_headers_footers(blocks, total_pages=6)

    # Should have removed the 6 header blocks and 6 footer blocks, kept the 6 content blocks
    assert len(stripped) == 6
    for block in stripped:
        assert "Real content" in block["content"]
        assert "Request for Proposal" not in block["content"]
        assert "Page " not in block["content"]


@pytest.mark.integration
def test_header_footer_stripper_preserves_unique_blocks():
    """Test that header_footer_stripper preserves non-repeated blocks."""
    # Unique headers on each page - should not be stripped
    # Use realistic page dimensions
    # Footer is NOT a simple page number (which would be stripped), but unique content
    blocks = [
        {"type": "text", "page": 1, "content": "Unique header on page 1", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 1, "content": "Content page 1", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 1, "content": "Footer note: Confidential", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 2, "content": "Different header page 2", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 2, "content": "Content page 2", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 2, "content": "Footer note: Internal use only", "bbox": [0, 760, 612, 792]},
        {"type": "text", "page": 3, "content": "Another unique header", "bbox": [0, 0, 612, 30]},
        {"type": "text", "page": 3, "content": "Content page 3", "bbox": [0, 150, 612, 200]},
        {"type": "text", "page": 3, "content": "Footer note: Draft version", "bbox": [0, 760, 612, 792]},
    ]

    stripped = strip_headers_footers(blocks, total_pages=3)

    # All blocks should be preserved since no block repeats on 60%+ of pages
    # (each header is unique, each footer is unique)
    assert len(stripped) == 9


@pytest.mark.integration
def test_pdf_extraction_no_ocr_needed():
    """Verify the fixture is a text-based PDF (not scanned) so OCR is not needed."""
    _check_fixture()

    with open(FIXTURE_PATH, "rb") as f:
        file_data = f.read()

    blocks = extract_pdf_content(file_data)

    # Text-based PDFs should extract substantial text content
    text_blocks = [b for b in blocks if b["type"] in ("text", "heading")]
    total_text = sum(len(b["content"]) for b in text_blocks)

    # Should have significant text content (not just a few OCR artifacts)
    assert total_text > 10000, f"Expected text-based PDF with substantial content, got {total_text} chars"


# DOCX test - create a simple test DOCX in memory
def _create_test_docx() -> bytes:
    """Create a simple test DOCX with heading and table."""
    import io

    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph with some content.")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("More content here.")

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[0].cells[2].text = "Header 3"
    table.rows[1].cells[0].text = "Row 1 Col 1"
    table.rows[1].cells[1].text = "Row 1 Col 2"
    table.rows[1].cells[2].text = "Row 1 Col 3"
    table.rows[2].cells[0].text = "Row 2 Col 1"
    table.rows[2].cells[1].text = "Row 2 Col 2"
    table.rows[2].cells[2].text = "Row 2 Col 3"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


@pytest.mark.integration
def test_docx_extraction():
    """Test DOCX extraction with heading and table."""
    file_data = _create_test_docx()

    blocks = extract_docx_content(file_data)

    # Should have heading, text, and table blocks
    block_types = [b["type"] for b in blocks]
    assert "heading" in block_types
    assert "text" in block_types
    assert "table" in block_types

    # Check table content
    table_blocks = [b for b in blocks if b["type"] == "table"]
    assert len(table_blocks) == 1
    table = table_blocks[0]["content"]
    assert len(table) == 3  # 3 rows
    assert len(table[0]) == 3  # 3 columns
    assert table[0][0] == "Header 1"
    assert table[1][0] == "Row 1 Col 1"